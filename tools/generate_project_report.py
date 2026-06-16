from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "TATA_Motors_Operations_Portal_Project_Report.docx"
ASSET_DIR = ROOT / "report_assets"
ASSET_DIR.mkdir(exist_ok=True)

USER_SCREENSHOTS = {
    "data_store": Path("/Users/sanyamsahu/Desktop/Screenshot 2026-06-15 at 14.04.51.png"),
    "loss_block": Path("/Users/sanyamsahu/Desktop/Screenshot 2026-06-15 at 14.05.00.png"),
    "meeting_sheet": Path("/Users/sanyamsahu/Desktop/Screenshot 2026-06-15 at 14.05.11.png"),
}


BLUE = "2E74B5"
DARK = "0B2545"
MUTED = "5B6472"
LIGHT = "F2F4F7"
BORDER = "CAD5E2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def style_doc(doc):
    section = doc.sections[0]
    configure_section(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 12, 6),
        ("Heading 2", 13, BLUE, 8, 4),
        ("Heading 3", 12, "1F4D78", 6, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def new_page_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)
    return section


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TATA MOTORS Operations Portal")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(DARK)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project Implementation Report")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("React + Google Apps Script + Google Sheets | 15 June 2026")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EEF6FF")
    set_cell_border(cell, "B8D7F4")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_workflow_table(doc):
    data = [
        ("1", "Shop representative", "Logs in and enters shift-wise production, downtime, failure count, hourly actuals, and loss details."),
        ("2", "React frontend", "Validates shift/date rules, calculates capacity, net downtime, OE, line efficiency, MTTR, MTBF, and performs optimistic local update."),
        ("3", "Apps Script API", "Receives saveReport request, locks the script, saves or updates raw rows, and returns confirmed JSON."),
        ("4", "Google Sheets", "Keeps Main Data as the hidden database and rebuilds Meeting Report, Loss Report, and Analytical Sheet as visible outputs."),
        ("5", "Admin review", "Admin views dashboards, exports workbook, checks missing uploads, and receives reminders/daily Excel reports."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(0.45), Inches(1.45), Inches(4.45)]
    headers = ["Step", "Stage", "What happens"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        cell.width = widths[idx]
        set_cell_shading(cell, LIGHT)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DARK)
    for row in data:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].width = widths[idx]
            set_cell_border(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[idx].paragraphs[0].add_run(text)


def snippet(path, start, end):
    lines = (ROOT / path).read_text(encoding="utf-8").splitlines()
    return [(i, lines[i - 1]) for i in range(start, end + 1)]


def make_code_image(filename, title, lines):
    font_path = Path("/System/Library/Fonts/SFNSMono.ttf")
    font = ImageFont.truetype(str(font_path), 18)
    title_font = ImageFont.truetype(str(font_path), 22)
    line_h = 25
    gutter = 72
    pad = 22
    max_len = max(len(text.expandtabs(2)) for _, text in lines)
    width = min(1900, max(1500, gutter + pad * 2 + max_len * 11))
    height = pad * 2 + 36 + len(lines) * line_h
    img = Image.new("RGB", (width, height), "#0B1220")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle([10, 10, width - 10, height - 10], radius=18, fill="#111827", outline="#334155", width=2)
    draw.text((pad, pad), title, fill="#E5E7EB", font=title_font)
    y = pad + 42
    for number, text in lines:
        draw.text((pad, y), f"{number:>4}", fill="#93A4B8", font=font)
        wrapped = text.expandtabs(2)[:145]
        color = "#D1D5DB"
        stripped = wrapped.strip()
        if stripped.startswith(("function", "async function", "const ", "let ", "var ")):
            color = "#BAE6FD"
        elif stripped.startswith(("if", "return", "try", "catch", "for")):
            color = "#C4B5FD"
        elif stripped.startswith("//"):
            color = "#86EFAC"
        draw.text((gutter + pad, y), wrapped, fill=color, font=font)
        y += line_h
    out = ASSET_DIR / filename
    img.save(out)
    return out


def add_image(doc, path, caption, width=6.3, page_break_before=False):
    p = doc.add_paragraph()
    if page_break_before:
        p.paragraph_format.page_break_before = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    cap.paragraph_format.space_after = Pt(6)


def crop_image(path, out_name, box):
    img = Image.open(path)
    cropped = img.crop(box)
    out = ASSET_DIR / out_name
    cropped.save(out)
    return out


def copy_image(path, out_name):
    img = Image.open(path).convert("RGB")
    out = ASSET_DIR / out_name
    img.save(out, quality=92)
    return out


def build():
    code_calc = make_code_image(
        "code_frontend_kpi.png",
        "Frontend KPI calculation and normalization",
        snippet("plantops-app/src/App.jsx", 415, 435),
    )
    code_submit = make_code_image(
        "code_submit_flow.png",
        "Submit flow: validation, optimistic update, Apps Script sync",
        snippet("plantops-app/src/App.jsx", 1080, 1110),
    )
    code_backend = make_code_image(
        "code_backend_sheets.png",
        "Apps Script API and Meeting Sheet rebuild",
        snippet("PlantOps_AppScript.gs", 294, 335) + snippet("PlantOps_AppScript.gs", 542, 565),
    )

    app_shot = crop_image(ROOT / "plantops_app_output.png", "output_input_form.png", (0, 0, 1280, 600))
    data_store_shot = copy_image(USER_SCREENSHOTS["data_store"], "output_data_store_actual.jpg")
    loss_block_shot = copy_image(USER_SCREENSHOTS["loss_block"], "output_loss_block_actual.jpg")
    meeting_sheet_shot = copy_image(USER_SCREENSHOTS["meeting_sheet"], "output_meeting_sheet_actual.jpg")

    doc = Document()
    style_doc(doc)
    add_title(doc)

    add_callout(
        doc,
        "Idea in one line: replace manual shop-wise production reporting with a web portal that collects shift data once, calculates KPIs automatically, and generates meeting-ready Google Sheet outputs.",
    )

    doc.add_heading("1. Project Idea and Need", level=1)
    p = doc.add_paragraph()
    p.add_run("Plant operations meetings need accurate shift-wise production, downtime, loss, and reliability data. ").bold = True
    p.add_run(
        "The existing workflow was Excel-oriented, which meant repeated manual entry, formula risk, and slower admin review. The project idea was to keep the familiarity of Excel/Google Sheets for managers, while giving shop representatives a guided digital form for daily submission."
    )
    add_bullets(
        doc,
        [
            "Shop representatives enter production, downtime, failure count, hourly actuals, and loss details.",
            "Fixed shop values such as cycle time, shift timing, QR, and targets are controlled in code.",
            "The backend converts raw submissions into Loss Report, Meeting Report, and Analytical Sheet outputs.",
        ],
    )

    doc.add_heading("2. How We Thought Through the Implementation", level=1)
    doc.add_paragraph(
        "The implementation follows the real plant-meeting workflow: collect data at the source, calculate KPIs consistently, preserve raw data, and present outputs in a format that operations teams can use directly. React supports fast forms and dashboards; Apps Script sits directly on Google Sheets and automates reports without a separate server."
    )
    add_workflow_table(doc)

    doc.add_heading("3. Technical Stack and Usage", level=1)
    doc.add_paragraph(
        "React + Vite is used for the portal, Recharts for analytics, lucide-react for icons, xlsx for workbook export, localStorage for cached settings, Google Apps Script for the API layer, and Google Sheets for storage plus visible meeting/analytical outputs."
    )

    new_page_section(doc)
    doc.add_heading("4. Main Code Screenshots", level=1)
    doc.add_paragraph(
        "The following screenshots show only the main implementation areas: KPI calculation, validated submission/sync, and backend sheet generation."
    )
    add_image(doc, code_calc, "Code screenshot 1: frontend calculation of capacity, downtime, OE, line efficiency, MTTR, and MTBF.", 5.85)
    add_image(doc, code_submit, "Code screenshot 2: submission flow with validation, optimistic UI update, and Google Sheets sync.", 5.85)

    new_page_section(doc)
    add_image(doc, code_backend, "Code screenshot 3: Apps Script endpoints and visible sheet rebuilding logic.", 5.75)

    new_page_section(doc)
    doc.add_heading("5. Output Screenshots", level=1)
    doc.add_paragraph(
        "The output screenshots show the working portal screens after data is available: the shop input form, the Data Store table, the block-style Loss and Meeting Data view, and the final Meeting Sheet export view."
    )
    add_image(doc, app_shot, "Output screenshot 1: shop data-entry form with fixed configuration and calculated actual time.", 5.6)
    add_image(doc, data_store_shot, "Output screenshot 2: Data Store view showing calculated production, downtime, AR, PR, QR, OE, and line efficiency columns.", 5.8)
    add_image(doc, loss_block_shot, "Output screenshot 3: Loss and Meeting Data single-entry block showing production and loss-time summary.", 5.8)
    add_image(doc, meeting_sheet_shot, "Output screenshot 4: Meeting Sheet view with production target, actual production, downtime, occurrence cards, and meeting table.", 5.8)

    doc.add_heading("6. Final Summary", level=1)
    add_bullets(
        doc,
        [
            "The project combines a structured web form with spreadsheet-native reporting, so shop users and admin users both get a familiar workflow.",
            "Calculations are centralized, which improves consistency across shop entries, dashboard views, and generated meeting reports.",
            "The hidden Main Data sheet works as the database, while rebuilt visible sheets serve as presentation-ready outputs.",
            "Automation through Apps Script triggers supports missing-upload reminders and daily Excel report emails.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("TATA MOTORS Operations Portal | Project Report")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
